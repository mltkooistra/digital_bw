import streamlit as st
import pandas as pd
import requests
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile
import os
from datetime import date
import statistics

# --- Pagina-instellingen ---
st.set_page_config(page_title="Genereer Rapport", layout="wide")
st.title("📄 Download groepsrapport")
st.text('Bedankt voor het meedoen aan de werksessie. In het document vind je een overzicht van de effecten op brede welvaart')

doc = Document("template_spg.docx")


# --- Sessieverificatie ---
if "access_code" not in st.session_state:
    st.error("Toegangscode ontbreekt.")
    st.stop()

if "group_answers_submitted" not in st.session_state or not st.session_state["group_answers_submitted"]:
    st.warning("De groepsantwoorden zijn nog niet ingediend. Ga terug en vul eerst de groepsvragen in.")
    st.stop()

# --- Gegevens ophalen ---
@st.cache_data(ttl=30)
def load_data():
    headers = {
        "apikey": st.secrets["supabase_key"],
        "Authorization": f"Bearer {st.secrets['supabase_key']}"
    }

    r_sub = requests.get(f"{st.secrets['supabase_url']}/rest/v1/submissions?select=*", headers=headers)
    df_sub = pd.DataFrame(r_sub.json()) if r_sub.status_code == 200 else pd.DataFrame()

    r_group = requests.get(f"{st.secrets['supabase_url']}/rest/v1/group_results?select=*", headers=headers)
    df_group = pd.DataFrame(r_group.json()) if r_group.status_code == 200 else pd.DataFrame()

    return df_sub, df_group

df_sub, df_group = load_data()
df_sub = df_sub[df_sub["session"] == st.session_state.access_code]
df_group = df_group[df_group["session"] == st.session_state.access_code]

if df_sub.empty or df_group.empty:
    st.warning("Niet genoeg data om een rapport te maken.")
    st.stop()

# --- Statistieken ---
n_participants = df_sub["name"].nunique()
n_groups = df_group["group"].nunique()

# --- Domeinen ---
domains = [
    "Welzijn", "Materiële welvaart", "Gezondheid", "Arbeid en vrije tijd",
    "Wonen", "Sociaal", "Veiligheid", "Milieu"
]

df_sub["signed_score"] = df_sub["score"] * df_sub["posneg"]
grouped = df_sub.groupby("domain")["signed_score"].mean().reindex(domains, fill_value=0)

# --- Diagram ---
def create_spider_chart(data):
    fig = go.Figure()
    fig.add_trace(go.Barpolar(
        r=[abs(v) for v in data],
        theta=domains,
        marker_color=["blue" if v >= 0 else "orange" for v in data],
        opacity=0.85
    ))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True)), showlegend=False)
    return fig

def save_plotly_chart(fig):
    img = BytesIO()
    fig.write_image(img, format="png")
    img.seek(0)
    return img

# --- Wordcloud functie ---
def generate_wordcloud(text):
    wc = WordCloud(width=800, height=400, background_color="white").generate(text)
    buf = BytesIO()
    plt.figure(figsize=(10, 5))
    plt.imshow(wc, interpolation="bilinear")
    plt.axis("off")
    plt.tight_layout()
    plt.savefig(buf, format="png")
    plt.close()
    buf.seek(0)
    return buf

# --- Hulpfuncties ---
def safe_int(val):
    try:
        return int(val)
    except (ValueError, TypeError):
        return None

def format_stats(values):
    if not values:
        return "geen data"
    return f"min: {min(values)} jaar, max: {max(values)} jaar, gemiddeld: {round(statistics.mean(values), 1)} jaar"

# --- Init document ---
doc = Document()
doc.add_heading(f"Verslag werksessie – {st.session_state.get('description', '–')}", 0)
doc.add_paragraph(f"Datum: {date.today().strftime('%d-%m-%Y')}")
doc.add_paragraph(f"Thema: {st.session_state.get('description', '–')}")
doc.add_paragraph(f"Informatie: {st.session_state.get('info', '–')}")
doc.add_paragraph(f"Aantal deelnemers: {n_participants}")
doc.add_paragraph(f"Aantal groepen: {n_groups}")
doc.add_page_break()

# --- Scores sectie ---
doc.add_heading("1. Gemiddelde scores per domein", level=1)
doc.add_paragraph("In onderstaande grafiek zie je hoe positief of negatief elk domein is beoordeeld door de deelnemers. Blauwe balken zijn positief, oranje negatief.")
doc.add_picture(save_plotly_chart(create_spider_chart(grouped)), width=Inches(6))
doc.add_page_break()

# --- Top effecten ---
df_group["votes"] = df_group.groupby("text")["text"].transform("count")
df_pos = df_group[df_group["votes"] > 0].sort_values("votes", ascending=False)
df_neg = df_group[df_group["votes"] < 0].sort_values("votes")

doc.add_heading("2. Hoogst gewaardeerde effecten", level=1)
top_n = n_groups * 3
for label, group_df in [("Positief", df_pos), ("Negatief", df_neg)]:
    doc.add_heading(f"{label} – meest genoemde effecten", level=2)
    for _, row in group_df.head(top_n).iterrows():
        doc.add_paragraph(f"• {row['text']} ({row['votes']} stemmen)")
doc.add_page_break()

# --- Samenvatting ---
pos_groups, neg_groups = [], []
pos_places, neg_places = [], []
pos_reach, neg_reach = [], []

pos_start_vals = [safe_int(r.get("feedback_start")) for r in df_group[df_group["votes"] >= 0].to_dict(orient="records")]
neg_start_vals = [safe_int(r.get("feedback_start")) for r in df_group[df_group["votes"] < 0].to_dict(orient="records")]
pos_start_vals = [v for v in pos_start_vals if v is not None]
neg_start_vals = [v for v in neg_start_vals if v is not None]

doc.add_heading("3. Samenvatting wie waar wanneer", level=1)
doc.add_paragraph("Hier zie je hoe de positieve en negatieve effecten geconcentreerd zijn bij groepen, plekken of in de tijd")

doc.add_heading("Positieve effecten", level=2)
doc.add_paragraph(f"• Groepen: {', '.join(filter(None, pos_groups))}")
doc.add_paragraph(f"• Plaatsen: {', '.join(filter(None, pos_places))}")
doc.add_paragraph(f"• Reikwijdte: {', '.join(filter(None, pos_reach))}")
doc.add_paragraph(f"• Verwachte start effect: {format_stats(pos_start_vals)}")

doc.add_heading("Negatieve effecten", level=2)
doc.add_paragraph(f"• Groepen: {', '.join(filter(None, neg_groups))}")
doc.add_paragraph(f"• Plaatsen: {', '.join(filter(None, neg_places))}")
doc.add_paragraph(f"• Reikwijdte: {', '.join(filter(None, neg_reach))}")
doc.add_paragraph(f"• Verwachte start effect: {format_stats(neg_start_vals)}")


doc.add_page_break()


# --- Details per effect ---
doc.add_heading("4. Groepsfeedback voor de belangrijkste effecten", level=1)
for label, group_df in [("Positief", df_pos), ("Negatief", df_neg)]:
    doc.add_heading(f"{label}e effecten", level=2)
    for _, row in group_df.iterrows():
        doc.add_heading(f"Effect: {row['text']}", level=3)
        doc.add_paragraph(f"Groep: {row.get('group', '–')}")
        doc.add_paragraph(f"- Groepsimpact: {row.get('feedback_group_impact', '')}")
        doc.add_paragraph(f"- Plaatsimpact: {row.get('feedback_place_impact', '')}")
        doc.add_paragraph(f"- Reikwijdte: {row.get('feedback_distance', '')}")
        doc.add_paragraph(f"- Verbeteringen: {row.get('feedback_improvements', '')}")

        if label == "Positief":
            pos_groups.append(row.get('feedback_group_impact', ''))
            pos_places.append(row.get('feedback_place_impact', ''))
            pos_reach.append(row.get('feedback_distance', ''))
        else:
            neg_groups.append(row.get('feedback_group_impact', ''))
            neg_places.append(row.get('feedback_place_impact', ''))
            neg_reach.append(row.get('feedback_distance', ''))


# --- Thema-analyse ---
# --- Thema-analyse ---
doc.add_heading("5. Thema-analyse", level=1)
for domain in domains:
    doc.add_heading(domain, level=2)
    domain_df = df_sub[df_sub["domain"] == domain]
    doc.add_paragraph(f"Aantal stemmen in dit domein: {len(domain_df)}")

    # Genereer alleen een wordcloud als er geldige tekst is
    text = " ".join(domain_df["text"].astype(str)).strip()
    if text:
        wc_img = generate_wordcloud(text)
        doc.add_picture(wc_img, width=Inches(5.5))
    else:
        doc.add_paragraph("⚠️ Geen tekst beschikbaar voor dit domein.")



# --- Opslaan ---# --- Opslaan ---
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    tmp_path = tmp_file.name  # Store path early
    doc.save(tmp_path)

with open(tmp_path, "rb") as f:
    st.download_button(
        label="📄 Download rapport als Word-bestand",
        data=f.read(),  # Read data first
        file_name="groepsrapport.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Now that everything is done, try removing it safely
try:
    os.remove(tmp_path)
except PermissionError:
    st.warning("Het bestand kon niet worden verwijderd – mogelijk nog in gebruik.")
