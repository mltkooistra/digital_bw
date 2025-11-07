import streamlit as st
import pandas as pd
import requests
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile
import os
from datetime import date
import statistics
from pathlib import Path
import nltk

# --- Page setup ---
st.set_page_config(page_title="Genereer Rapport", layout="wide")
st.title("📄 Download groepsrapport")
st.text("Bedankt voor het meedoen aan de werksessie. In het document vind je een overzicht van de effecten op brede welvaart")

# ===============================
# NLTK Dutch stopwords (local dir)
# ===============================
NLTK_DIR = Path(__file__).resolve().parents[1] / ".nltk_data"
NLTK_DIR.mkdir(exist_ok=True)
if str(NLTK_DIR) not in nltk.data.path:
    nltk.data.path.insert(0, str(NLTK_DIR))

@st.cache_resource
def ensure_dutch_stopwords():
    # Make sure 'stopwords' corpus exists; download to our local folder if missing
    try:
        nltk.data.find("corpora/stopwords")
    except LookupError:
        nltk.download("stopwords", download_dir=str(NLTK_DIR))
        if str(NLTK_DIR) not in nltk.data.path:
            nltk.data.path.insert(0, str(NLTK_DIR))

    from nltk.corpus import stopwords
    try:
        return set(stopwords.words("dutch"))
    except OSError:
        # Fallback (offline/no download): minimal set so the app keeps working
        return {
            "de","het","een","en","of","maar","want","dat","die","dit","er","je","jij",
            "u","we","wij","ze","zij","ik","hij","in","op","aan","met","voor",
            "van","naar","bij","als","dan","niet","geen","wel","ook","om","te","tot",
        }

dutch_stopwords = ensure_dutch_stopwords()

# --- Session checks ---
if "access_code" not in st.session_state:
    st.error("Toegangscode ontbreekt.")
    st.stop()

SESSION = st.session_state.access_code

# --- Data loading ---
@st.cache_data(ttl=30)
def load_data(session_code: str):
    headers = {
        "apikey": st.secrets["supabase_key"],
        "Authorization": f"Bearer {st.secrets['supabase_key']}"
    }

    # Submissions (alle individuele inzendingen)
    r_sub = requests.get(f"{st.secrets['supabase_url']}/rest/v1/submissions?select=*", headers=headers)
    df_sub = pd.DataFrame(r_sub.json()) if r_sub.status_code == 200 else pd.DataFrame()

    # Group results (ingediende groepsopdracht)
    r_group = requests.get(f"{st.secrets['supabase_url']}/rest/v1/group_results?select=*", headers=headers)
    df_group = pd.DataFrame(r_group.json()) if r_group.status_code == 200 else pd.DataFrame()

    # Groups (wie zitten er in welke groep)
    r_groups = requests.get(
        f"{st.secrets['supabase_url']}/rest/v1/groups?select=group,session",
        headers=headers
    )
    df_groups_all = pd.DataFrame(r_groups.json()) if r_groups.status_code == 200 else pd.DataFrame()

    return df_sub, df_group, df_groups_all

df_sub, df_group, df_groups_all = load_data(SESSION)

# Filter op sessie
df_sub = df_sub[df_sub.get("session") == SESSION].copy()
df_group = df_group[df_group.get("session") == SESSION].copy()
df_groups_all = df_groups_all[df_groups_all.get("session") == SESSION].copy()

if df_sub.empty:
    st.warning("Er zijn nog geen individuele inzendingen voor deze sessie.")
    st.stop()

# --- Determine completion status per group ---
# Groepen die bestaan in deze sessie:
groups_known = sorted(
    g for g in df_groups_all.get("group", pd.Series(dtype=str)).dropna().astype(str).unique()
)
# Fallback: als tabel 'groups' leeg is, probeer uit group_results
if not groups_known:
    groups_known = sorted(
        g for g in df_group.get("group", pd.Series(dtype=str)).dropna().astype(str).unique()
    )

# Groepen die iets hebben ingediend in group_results
groups_submitted = sorted(
    g for g in df_group.get("group", pd.Series(dtype=str)).dropna().astype(str).unique()
)

groups_missing = [g for g in groups_known if g not in groups_submitted]
all_done = (len(groups_known) > 0) and (len(groups_missing) == 0)

# --- Status banner ---
left, right = st.columns([0.7, 0.3])
with left:
    if all_done:
        st.success(f"✅ Alle {len(groups_known)} groepen hebben de groepsopdracht ingediend. Het rapport is compleet.")
    else:
        if groups_known:
            st.warning(
                f"⚠️ Rapport is **nog niet compleet**: "
                f"{len(groups_submitted)}/{len(groups_known)} groepen ingediend.\n\n"
                f"Ontbrekend(e): {', '.join(groups_missing) if groups_missing else '—'}"
            )
        else:
            st.warning("⚠️ Geen groepen gevonden voor deze sessie. Het rapport kan onvolledig zijn.")

with right:
    st.metric("Ingediend", f"{len(groups_submitted)} / {max(1, len(groups_known))}")

# --- Stats ---
n_participants = df_sub.get("name", pd.Series(dtype=str)).nunique()
n_groups = len(groups_known) if groups_known else df_group.get("group", pd.Series(dtype=str)).nunique()

# --- Domains ---
domains = [
    "Welzijn", "Materiële welvaart", "Gezondheid", "Arbeid en vrije tijd",
    "Wonen", "Sociaal", "Veiligheid", "Milieu"
]

# Zorg dat kolommen bestaan
for col in ["score", "posneg", "domain", "text"]:
    if col not in df_sub.columns:
        df_sub[col] = 0 if col in ("score","posneg") else ""

# Signed score (posneg ∈ {1,-1})
df_sub["score"] = pd.to_numeric(df_sub["score"], errors="coerce").fillna(0)
df_sub["posneg"] = pd.to_numeric(df_sub["posneg"], errors="coerce").fillna(0)
df_sub["signed_score"] = df_sub["score"] * df_sub["posneg"]
grouped = (
    df_sub.groupby("domain")["signed_score"].mean()
    .reindex(domains, fill_value=0).tolist()
)

# --- Chart helpers ---
def create_spider_chart(data):
    fig = go.Figure()
    fig.add_trace(go.Barpolar(
        r=[abs(v) for v in data],
        theta=domains,
        marker_color=["blue" if v >= 0 else "orange" for v in data],
        opacity=0.85
    ))
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True)),
        showlegend=False,
        margin=dict(l=0, r=0, t=0, b=0)
    )
    return fig

def save_plotly_chart(fig):
    """
    Try exporting with Plotly (kaleido). If that fails (e.g., kaleido missing),
    fall back to a simple matplotlib bar chart so the DOCX still gets an image.
    """
    img = BytesIO()
    try:
        fig.write_image(img, format="png")
        img.seek(0)
        return img
    except Exception:
        # Fallback: matplotlib horizontal bar of signed scores
        plt.figure(figsize=(7, 4))
        vals = grouped  # use already computed list
        colors = ["tab:blue" if v >= 0 else "tab:orange" for v in vals]
        plt.barh(list(domains), vals, alpha=0.85, color=colors)
        plt.axvline(0, linewidth=1)
        plt.tight_layout()
        plt.savefig(img, format="png")
        plt.close()
        img.seek(0)
        return img

# --- Wordcloud (uses Dutch stopwords) ---
def generate_wordcloud(text):
    wc = WordCloud(
        width=800,
        height=400,
        background_color="white",
        stopwords=dutch_stopwords
    ).generate(text)
    buf = BytesIO()
    plt.figure(figsize=(10, 5))
    plt.imshow(wc, interpolation="bilinear")
    plt.axis("off")
    plt.tight_layout()
    plt.savefig(buf, format="png")
    plt.close()
    buf.seek(0)
    return buf

# --- Utilities ---
def safe_int(val):
    try:
        return int(val)
    except (ValueError, TypeError):
        return None

def format_stats(values):
    if not values:
        return "geen data"
    return f"min: {min(values)} jaar, max: {max(values)} jaar, gemiddeld: {round(statistics.mean(values), 1)} jaar"

# =========================
# Ensure expected group_results columns (NEW QUESTIONS)
# =========================
for col, default in [
    ("direction", None),
    ("feedback_group_impact", ""),
    ("feedback_distance", []),      # JSON array (list)
    ("feedback_start", ""),         # categorical: direct/weken/...
    ("feedback_improvements", ""),
    ("group", ""),
    ("text", ""),
    ("domein", ""),
]:
    if col not in df_group.columns:
        df_group[col] = default

# Normalize JSON/list-like columns that may be strings in DB
def _as_list(v):
    if isinstance(v, list):
        return v
    if isinstance(v, str) and v.strip():
        try:
            import json as _json
            parsed = _json.loads(v)
            if isinstance(parsed, list):
                return parsed
        except Exception:
            return [p.strip() for p in v.split(",") if p.strip()]
    return []

df_group["feedback_distance"] = df_group["feedback_distance"].apply(_as_list)

# Split pos/neg datasets
df_pos = df_group[df_group["direction"].astype(str).str.lower().eq("pos")].copy()
df_neg = df_group[df_group["direction"].astype(str).str.lower().eq("neg")].copy()

# Derive a simple "votes" count per effect text within each polarity (if no explicit votes column)
if "votes" not in df_pos.columns:
    df_pos["votes"] = df_pos.groupby("text")["text"].transform("count")
if "votes" not in df_neg.columns:
    df_neg["votes"] = df_neg.groupby("text")["text"].transform("count")

# Sort by votes
df_pos = df_pos.sort_values("votes", ascending=False)
df_neg = df_neg.sort_values("votes", ascending=False)

# --- Build the DOCX in memory (always build so both buttons can use it) ---
doc = Document()
doc.add_heading(f"Verslag werksessie – {st.session_state.get('description', '–')}", 0)
doc.add_paragraph(f"Datum: {date.today().strftime('%d-%m-%Y')}")
doc.add_paragraph(f"Thema: {st.session_state.get('description', '–')}")
doc.add_paragraph(f"Informatie: {st.session_state.get('info', '–')}")
doc.add_paragraph(f"Aantal deelnemers: {n_participants}")
doc.add_paragraph(f"Aantal groepen: {n_groups}")
if not all_done:
    doc.add_paragraph("⚠️ Let op: dit rapport is mogelijk onvolledig. Niet alle groepen hebben hun groepsopdracht ingediend.")
doc.add_page_break()

# --- Scores section ---
doc.add_heading("1. Gemiddelde scores per domein", level=1)
doc.add_paragraph("In onderstaande grafiek zie je hoe positief of negatief elk domein is beoordeeld door de deelnemers. Blauwe balken zijn positief, oranje negatief.")
doc.add_picture(save_plotly_chart(create_spider_chart(grouped)), width=Inches(6))
doc.add_page_break()

# --- Top effects (uses pos/neg splits) ---
doc.add_heading("2. Hoogst gewaardeerde effecten", level=1)
top_n = max(1, (n_groups if n_groups else 1) * 3)

for label, group_df in [("Positief", df_pos), ("Negatief", df_neg)]:
    doc.add_heading(f"{label} – meest genoemde effecten", level=2)
    if group_df.empty:
        doc.add_paragraph("— Geen effecten —")
        continue
    for _, row in group_df.head(top_n).iterrows():
        doc.add_paragraph(f"• {row['text']} ({row['votes']} stemmen)")
doc.add_page_break()

# --- Summary (no 'plaatsen' anymore) ---
def summarize_reach(list_series):
    # Flatten list-of-lists and count top items
    from collections import Counter
    all_items = []
    for v in list_series:
        if isinstance(v, list):
            all_items.extend([str(x) for x in v if str(x)])
    if not all_items:
        return "—"
    cnt = Counter(all_items)
    # show top 5: item (count)
    top = ", ".join([f"{k} ({v})" for k, v in cnt.most_common(5)])
    return top

def summarize_start(cat_series):
    vals = [str(v) for v in cat_series if isinstance(v, str) and v.strip()]
    if not vals:
        return "—"
    from collections import Counter
    cnt = Counter(vals)
    mode, mode_n = cnt.most_common(1)[0]
    total = sum(cnt.values())
    parts = [f"{k}: {v}" for k, v in cnt.most_common()]
    return f"meest genoemd: {mode} ({mode_n}/{total}); verdeling: " + ", ".join(parts)

doc.add_heading("3. Samenvatting wie – reikwijdte – wanneer", level=1)
doc.add_paragraph("Overzicht per polariteit (positieve en negatieve effecten).")

# Positive
doc.add_heading("Positieve effecten", level=2)
doc.add_paragraph(
    "• Voor wie is het effect het grootst? " +
    (', '.join([s for s in df_pos['feedback_group_impact'].astype(str).tolist() if s]) or '—')
)
doc.add_paragraph(f"• Reikwijdte (meest genoemd): {summarize_reach(df_pos['feedback_distance'].tolist())}")
doc.add_paragraph(f"• Verwachte start (samenvatting): {summarize_start(df_pos['feedback_start'].tolist())}")

# Negative
doc.add_heading("Negatieve effecten", level=2)
doc.add_paragraph(
    "• Voor wie is het effect het grootst? " +
    (', '.join([s for s in df_neg['feedback_group_impact'].astype(str).tolist() if s]) or '—')
)
doc.add_paragraph(f"• Reikwijdte (meest genoemd): {summarize_reach(df_neg['feedback_distance'].tolist())}")
doc.add_paragraph(f"• Verwachte start (samenvatting): {summarize_start(df_neg['feedback_start'].tolist())}")
doc.add_page_break()

# --- Details per effect (no 'plaatsimpact', conditional last question label) ---
doc.add_heading("4. Groepsfeedback voor de belangrijkste effecten", level=1)

def reach_as_text(v):
    if isinstance(v, list):
        return ", ".join([str(x) for x in v if str(x)]) or "—"
    return str(v) if str(v) else "—"

for label, group_df in [("Positief", df_pos), ("Negatief", df_neg)]:
    doc.add_heading(f"{label}e effecten", level=2)
    if group_df.empty:
        doc.add_paragraph("— Geen effecten —")
        continue

    for _, row in group_df.iterrows():
        doc.add_heading(f"Effect: {row.get('text','')}", level=3)
        doc.add_paragraph(f"Groep: {row.get('group', '–')}")
        doc.add_paragraph(f"- Voor wie is dit effect het grootst? {row.get('feedback_group_impact', '') or '—'}")
        doc.add_paragraph(f"- Reikwijdte: {reach_as_text(row.get('feedback_distance', []))}")
        doc.add_paragraph(f"- Wanneer verwacht: {row.get('feedback_start', '') or '—'}")

        # Conditional improvements phrasing
        if label.startswith("Pos"):
            improv_label = "Zijn er aanpassingen mogelijk om het effect te versterken?"
        else:
            improv_label = "Zijn er aanpassingen mogelijk aan de interventie om dit effect te beperken of voorkomen?"

        doc.add_paragraph(f"- {improv_label} {row.get('feedback_improvements', '') or '—'}")

# --- Theme analysis (with Dutch stopwords in wordcloud) ---
doc.add_heading("5. Thema-analyse", level=1)
for domain in domains:
    doc.add_heading(domain, level=2)
    domain_df = df_sub[df_sub["domain"] == domain]
    doc.add_paragraph(f"Aantal stemmen in dit domein: {len(domain_df)}")

    text = " ".join(domain_df["text"].astype(str)).strip()
    if text:
        wc_img = generate_wordcloud(text)
        doc.add_picture(wc_img, width=Inches(5.5))
    else:
        doc.add_paragraph("⚠️ Geen tekst beschikbaar voor dit domein.")

# --- Save to a temp file once (used by both download buttons) ---
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    tmp_path = tmp_file.name
    doc.save(tmp_path)

# --- Download buttons ---
colA, colB = st.columns(2)

with open(tmp_path, "rb") as f1:
    with colA:
        st.download_button(
            label="✅ Download volledig rapport",
            data=f1.read(),
            file_name="groepsrapport.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not all_done,
            help=None if all_done else "Niet alle groepen hebben ingediend; deze knop wordt actief zodra iedereen klaar is."
        )

with open(tmp_path, "rb") as f2:
    with colB:
        st.download_button(
            label="⚠️ Download toch (onvolledig)",
            data=f2.read(),
            file_name="groepsrapport_onvolledig.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Download het rapport zoals het nu is, ook als niet alle groepen hebben ingediend."
        )

# Cleanup
try:
    os.remove(tmp_path)
except PermissionError:
    st.warning("Het tijdelijke bestand kon niet worden verwijderd – mogelijk nog in gebruik.")
