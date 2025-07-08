import streamlit as st
import pandas as pd
import requests
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile
import os

# --- Setup ---
st.set_page_config(page_title="Genereer Rapport", layout="wide")
st.title("📄 Download groepsrapport")

# --- Check sessie ---
if "access_code" not in st.session_state:
    st.error("Toegangscode ontbreekt.")
    st.stop()

# --- Laad gegevens ---
@st.cache_data(ttl=30)
def load_data():
    headers = {
        "apikey": st.secrets["supabase_key"],
        "Authorization": f"Bearer {st.secrets['supabase_key']}"
    }

    # Submissions
    r_sub = requests.get(
        f"{st.secrets['supabase_url']}/rest/v1/submissions?select=*",
        headers=headers
    )
    df_sub = pd.DataFrame(r_sub.json()) if r_sub.status_code == 200 else pd.DataFrame()

    # Group results
    r_group = requests.get(
        f"{st.secrets['supabase_url']}/rest/v1/group_results?select=*",
        headers=headers
    )
    df_group = pd.DataFrame(r_group.json()) if r_group.status_code == 200 else pd.DataFrame()

    return df_sub, df_group

df_sub, df_group = load_data()
df_sub = df_sub[df_sub["session"] == st.session_state.access_code]
df_group = df_group[df_group["session"] == st.session_state.access_code]

if df_sub.empty or df_group.empty:
    st.warning("Niet genoeg data om een rapport te maken.")
    st.stop()

# --- Maak spider chart ---
domains = [
    "Welzijn", "Materiële welvaart", "Gezondheid", "Arbeid en vrije tijd",
    "Wonen", "Sociaal", "Veiligheid", "Milieu"
]

df_sub["signed_score"] = df_sub["score"] * df_sub["posneg"]
grouped = df_sub.groupby("domain")["signed_score"].mean().reindex(domains, fill_value=0)

def create_spider_chart(data):
    fig = go.Figure()
    fig.add_trace(go.Barpolar(
        r=[abs(v) for v in data],
        theta=domains,
        marker_color=["blue" if v >= 0 else "orange" for v in data],
        opacity=0.85
    ))
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True)),
        showlegend=False
    )
    return fig

# Save chart to image
def save_plotly_chart(fig):
    img_bytes = fig.to_image(format="png")
    return BytesIO(img_bytes)

# --- Maak wordcloud ---
all_text = " ".join(df_sub["text"].dropna().astype(str))
wordcloud_img = BytesIO()
wc = WordCloud(width=800, height=400, background_color="white").generate(all_text)
plt.figure(figsize=(10, 5))
plt.imshow(wc, interpolation="bilinear")
plt.axis("off")
plt.tight_layout()
plt.savefig(wordcloud_img, format="png")
plt.close()

wordcloud_img.seek(0)

# --- Maak Word document ---
doc = Document()
doc.add_heading("Groepsrapport", 0)

doc.add_heading("1. Gemiddelde scores per domein", level=1)
spider_fig = create_spider_chart(grouped)
spider_img = save_plotly_chart(spider_fig)
doc.add_picture(spider_img, width=Inches(6))

doc.add_heading("2. Wordcloud van effecten", level=1)
doc.add_picture(wordcloud_img, width=Inches(6))

doc.add_heading("3. Groepsfeedback", level=1)

for _, row in df_group.iterrows():
    doc.add_heading(f"Effect: {row['text']}", level=2)
    doc.add_paragraph(f"Groep: {row.get('group', '–')}")
    doc.add_paragraph(f"- Groepsimpact: {row.get('feedback_group_impact', '')}")
    doc.add_paragraph(f"- Plaatsimpact: {row.get('feedback_place_impact', '')}")
    doc.add_paragraph(f"- Reikwijdte: {row.get('feedback_distance', '')}")
    doc.add_paragraph(f"- Verbeteringen: {row.get('feedback_improvements', '')}")
    doc.add_paragraph("")

# --- Save to BytesIO and offer download ---
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc.save(tmp_file.name)
    with open(tmp_file.name, "rb") as f:
        st.download_button(
            label="📄 Download rapport als Word-bestand",
            data=f,
            file_name="groepsrapport.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    os.unlink(tmp_file.name)

